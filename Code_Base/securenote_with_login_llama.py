import os
import time
import json
import shutil
import logging
import warnings
import tempfile
import datetime
from pathlib import Path
from uuid import uuid4
from io import BytesIO
import types
import sys
from uuid import uuid4
import uuid
import requests


import pandas as pd
from numpy import dot
from numpy.linalg import norm
import numpy as np
import fitz  # PyMuPDF
import pdfplumber
from PIL import Image
import streamlit as st
import streamlit.components.v1 as components
from tqdm import tqdm
import camelot
import json

# os.environ["HUGGINGFACE_HUB_TOKEN"] = "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"
# from huggingface_hub import login, whoami
# from huggingface_hub.utils import HfHubHTTPError
# login(os.getenv("HUGGINGFACE_HUB_TOKEN"))


from langchain.vectorstores import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.memory import ConversationBufferMemory
from langchain.schema import HumanMessage, AIMessage
from sentence_transformers import SentenceTransformer
from chromadb import PersistentClient
from chromadb.config import Settings

# from unstructured.partition.pdf import partition_pdf
from docx.image.image import Image as DocxImage
from docx import Document as DocxDocument


if "text_vector_db" not in st.session_state:
    st.session_state.text_vector_db = None

if "image_vector_db" not in st.session_state:
    st.session_state.image_vector_db = None

if "table_vector_db" not in st.session_state:
    st.session_state.table_vector_db = None
    

st.title("🤖 SecureNote – Saint-Gobain Technical Document Assistant")
st.set_page_config(page_title="SecureNote - Saint-Gobain", layout="wide")

# Define simple demo credentials
allowed_users = ["krishna","falcon","ethan"]
stored_password = {
    "krishna": "krishnapass",
    "falcon": "falconpass",
    "ethan": "ethanpass"
}

# place login check early
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False
if "user_id" not in st.session_state:
    st.session_state["user_id"] = None

if not st.session_state["authenticated"]:
    st.title("🔒 SecureNote Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        # you can replace this check with your own secure auth later
        if username in allowed_users and password == stored_password[username]:
            st.session_state["authenticated"] = True
            st.session_state["user_id"] = username
            st.rerun()
        else:
            st.error("❌ Invalid credentials. Try again.")
    st.stop()


if st.session_state.get("authenticated"):
# place logout button on top right with Streamlit
    col1, col2 = st.columns([11,1])
    with col2:
        if st.button("🚪 Logout"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()


if "new_uploads" not in st.session_state:
    st.session_state["new_uploads"] = []
    
# Initialize session state variables
if "uploader_key" not in st.session_state:
    st.session_state["uploader_key"] = 0
    
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("unstructured").setLevel(logging.ERROR)

dummy_module = types.ModuleType("torch.classes")
dummy_module.__path__ = []
sys.modules["torch.classes"] = dummy_module

BASE_DIR = Path.home() / "Documents" / "Projects" / "RAG_Chatbot" / "Code_base"

ROOT_FOLDER = BASE_DIR / "data_store" / st.session_state['user_id']
ROOT_FOLDER.mkdir(parents=True, exist_ok=True)

CHAT_FOLDER = ROOT_FOLDER / "chats"
CHAT_FOLDER.mkdir(exist_ok=True)


def list_md_files(folder):
    return list(Path(folder).glob("*.md"))


UPLOAD_FOLDER = ROOT_FOLDER / "uploaded"
UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)

NOTES_FOLDER = BASE_DIR / "Notes"

existing_md_files = list_md_files(NOTES_FOLDER)

text_db_path = ROOT_FOLDER / "chroma_text_db"
image_db_path = ROOT_FOLDER / "chroma_image_db"
table_db_path = ROOT_FOLDER / "chroma_table_db"

memory = ConversationBufferMemory(k=20, return_messages=True) # type: ignore

active_chat_file = None
active_chat_title = None

system_prompt = """This is AIRA, created by Saint-Gobain. AIRA specializes in analyzing and responding to queries about  technical documentation and reports.

AIRA's knowledge comes exclusively from the technical documents provided. AIRA's information was last updated with document set [document_set_id].

AIRA cannot verify information outside the provided documents. If asked about information not covered in the documents, AIRA lets the human know this directly and simply.

If presented with a technical problem or analysis request, AIRA thinks through it step by step before giving its final answer.


AIRA provides:
- Direct, concise answers for specific queries about facts, figures, or simple information
- Thorough responses for technical analysis, processes, and complex questions
- Clear citations to source documents

When AIRA mentions or cites specific documents, it always includes the document ID or reference number.

If AIRA finds inconsistencies between documents or unclear information, it points this out to the human.

AIRA uses markdown for formatting:
- Code blocks for technical specifications or formulas
- **Bold** for emphasis
- Lists for multiple points
- Headers for organizing longer responses

If there is ambiguity in the human's question, AIRA asks for clarification to provide the most accurate response.

AIRA proceeds directly with answering questions without asking permission first.

If AIRA cannot find information about a specific query in the documents, it simply states "This information is not found in the provided documents" rather than speculating.
"""





document_paths = (
    list(UPLOAD_FOLDER.glob("*.pdf")) +
    list(UPLOAD_FOLDER.glob("*.docx")) +   # ✅ Add this
    list(NOTES_FOLDER.glob("*.md"))
)


embedding_model = SentenceTransformer('sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2')

import requests
import streamlit as st

def query_llama(messages):
    prompt = ""
    for msg in messages:
        if msg["role"] == "system":
            prompt += f"<<SYS>>\n{msg['content']}\n<</SYS>>\n\n"
        elif msg["role"] == "user":
            prompt += f"[INST] {msg['content']} [/INST]\n"
        elif msg["role"] == "assistant":
            prompt += f"{msg['content']}\n"

    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama3",
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.1,
                    "top_p": 0.95,
                    "num_predict": 1024
                }
            },
            timeout=60
        )
        response.raise_for_status()
        data = response.json()
        # st.write(data)  # shows the *entire* JSON
        answer = data.get("response", "⚠️ No response from LLaMA.").strip()


        # clean up  
        answer = answer.replace("<<SYS>>", "")
        answer = answer.replace("<<ANSWER>>", "")
        answer = answer.replace("[INST]", "")
        answer = answer.replace("[/INST]", "")
        answer = answer.strip()

    
        return answer

    except requests.exceptions.RequestException as e:
        st.error(f"⚠️ Request to LLaMA server failed: {e}")
        return "⚠️ LLaMA server error."


def get_embedding(text):
    try:
        # model.encode returns a NumPy array
        return embedding_model.encode(text).tolist()
    except Exception as e:
        print(f"Embedding failed: {e}")
        return np.zeros(384).tolist()
    
    
def cosine_similarity(a, b):
    return dot(a, b) / (norm(a) * norm(b) + 1e-8)

    
    
def process_query(user_query):
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_query},
    ]
    answer = query_llama(messages)
    return answer


# Ensure DBs are reloaded after restart
def load_existing_dbs():
    # Load text DB
    if "text_vector_db" not in st.session_state or st.session_state.text_vector_db is None:
        if text_db_path.exists():
            chroma_client = PersistentClient(path=str(text_db_path), settings=Settings(anonymized_telemetry=False))
            st.session_state.text_vector_db = Chroma(
                client=chroma_client,
                collection_name="default",
                persist_directory=str(text_db_path)
            )

    # Load image DB
    if "image_vector_db" not in st.session_state or st.session_state.image_vector_db is None:
        if image_db_path.exists():
            chroma_client = PersistentClient(path=str(image_db_path), settings=Settings(anonymized_telemetry=False))
            st.session_state.image_vector_db = Chroma(
                client=chroma_client,
                collection_name="default",
                persist_directory=str(image_db_path)
            )

    # Load table DB
    if "table_vector_db" not in st.session_state or st.session_state.table_vector_db is None:
        if table_db_path.exists():
            chroma_client = PersistentClient(path=str(table_db_path), settings=Settings(anonymized_telemetry=False))
            st.session_state.table_vector_db = Chroma(
                client=chroma_client,
                collection_name="default",
                persist_directory=str(table_db_path)
            )

print("💾 Attempting to load existing DBs...")
load_existing_dbs()
print("✅ DB load completed.")

def save_chat_to_file(memory):
    global active_chat_file, active_chat_title
    if active_chat_file is None:
        print("⚠️ No active chat file. Chat not saved.")
        return
    messages = [{"role": m.type, "content": m.content} for m in memory.chat_memory.messages]
    with open(active_chat_file, 'w') as f:
        json.dump({"title": active_chat_title, "messages": messages}, f)
    print(f"✅ Chat saved to {active_chat_file.name}")

    # Keep only latest 10 files
    all_chats = sorted(CHAT_FOLDER.glob("chat_*.json"), reverse=True)
    max_files = 10
    for old_file in all_chats[max_files:]:
        old_file.unlink()
    
    
def list_previous_chats():
    chat_files = sorted(CHAT_FOLDER.glob("chat_*.json"), reverse=True)[:7]
    chats = []
    for idx, file in enumerate(chat_files, 1):
        with open(file) as f:
            data = json.load(f)
            title = data.get("title", file.stem.replace("chat_", ""))
            turn_count = len(data.get("messages", [])) // 2  # rough user-assistant pairs
            chats.append((idx, file, title, turn_count))
    return chats
    
def speak_text_button(text, key_suffix="default"):
    html_id_speak = f"speak_button_{key_suffix}"
    html_id_stop = f"stop_button_{key_suffix}"
    escaped_text = json.dumps(text)

    html_code = f"""
    <button id="{html_id_speak}" style="margin: 6px 6px 0 0; padding: 6px 14px; border-radius: 6px; border: none; background-color: #4CAF50; color: white; font-weight: bold; cursor: pointer;">
        🔊 Read Aloud
    </button>
    <button id="{html_id_stop}" style="margin-top: 6px; padding: 6px 14px; border-radius: 6px; border: none; background-color: #e53935; color: white; font-weight: bold; cursor: pointer;">
        🛑 Stop
    </button>
    <script>
    const btnSpeak = document.getElementById("{html_id_speak}");
    const btnStop = document.getElementById("{html_id_stop}");

    if (btnSpeak) {{
        btnSpeak.onclick = function() {{
            const msg = new SpeechSynthesisUtterance({escaped_text});
            msg.rate = 1;
            msg.pitch = 1;
            msg.lang = "en-US";
            window.speechSynthesis.cancel();
            window.speechSynthesis.speak(msg);
        }};
    }}

    if (btnStop) {{
        btnStop.onclick = function() {{
            window.speechSynthesis.cancel();
        }};
    }}
    </script>
    """
    components.html(html_code, height=60)


# def speak_text_button(text, key_suffix="default"):
#     html_id_speak = f"speak_button_{key_suffix}"
#     html_id_stop = f"stop_button_{key_suffix}"
#     escaped_text = json.dumps(text)

#     html_code = f"""
#     <button id="{html_id_speak}" style="margin: 6px 6px 0 0; padding: 6px 12px; border-radius: 5px; border: none; background-color: #4CAF50; color: white; font-weight: bold; cursor: pointer;">
#         🔊 Read Aloud
#     </button>
#     <button id="{html_id_stop}" style="margin-top: 6px; padding: 6px 12px; border-radius: 5px; border: none; background-color: #e53935; color: white; font-weight: bold; cursor: pointer;">
#         🛑 Stop
#     </button>
#     <script>
#     let speakMsg;

#     function speak() {{
#         const voices = window.speechSynthesis.getVoices();
#         if (!voices.length) {{
#             // Retry once voices are loaded
#             window.speechSynthesis.onvoiceschanged = () => {{
#                 speak();
#             }};
#             return;
#         }}

#         speakMsg = new SpeechSynthesisUtterance({escaped_text});
#         speakMsg.voice = voices.find(v => v.name.includes('Google US English')) || voices[0];
#         speakMsg.rate = 1;
#         speakMsg.pitch = 1;
#         speakMsg.lang = "en-US";
#         window.speechSynthesis.cancel();
#         window.speechSynthesis.speak(speakMsg);
#     }}

#     document.getElementById("{html_id_speak}").onclick = function() {{
#         speak();
#     }};

#     document.getElementById("{html_id_stop}").onclick = function() {{
#         window.speechSynthesis.cancel();
#     }};
#     </script>
#     """
    # components.html(html_code, height=60)

def stream_text_to_ui(text, delay=0.1):
    placeholder = st.empty()
    lines = text.split('\n')
    streamed = ""
    for line in lines:
        streamed += line + "\n"
        placeholder.markdown(streamed + "▌")  # Typing cursor effect
        time.sleep(delay)
    placeholder.markdown(streamed) 



def load_docs(file_paths):
    all_splits = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)

    for file_path in file_paths:
        if file_path.suffix.lower() == ".pdf":
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            print("✅ Loaded pages:", len(pages), "from", file_path)
            
        elif file_path.suffix.lower() == ".docx":
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            pages = [Document(page_content=full_text, metadata={"source": str(file_path)})]
            
        elif file_path.suffix.lower() == ".md":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            pages = [Document(page_content=content, metadata={"source": str(file_path)})]
        else:
            continue

        splits = text_splitter.split_documents(pages)
        for doc in splits:
            doc.metadata["source"] = str(file_path)
            doc.page_content = f"[SOURCE: {file_path.name}]\n\n" + doc.page_content
        all_splits.extend(splits)

    return all_splits



def create_db(splits):
    
    print("🔥 Inside create_db with", len(splits), "splits")
    
    start_time = time.time()
    
    texts = [doc.page_content for doc in splits]
    metadatas = [doc.metadata for doc in splits]

    embeddings_list = []
    valid_texts = []
    valid_metadatas = []

    for i, text in enumerate(texts):
        emb = get_embedding(text)
        if emb and any(emb):
            embeddings_list.append(emb)
            valid_texts.append(text)
            valid_metadatas.append(metadatas[i])
        else:
            print(f"⚠️ Skipping empty embedding for doc chunk {i}")

    if not embeddings_list:
        raise ValueError("❌ No valid embeddings were generated. Cannot create DB.")

    permanent_chroma_path_text = ROOT_FOLDER / "chroma_text_db"
    permanent_chroma_path_text.mkdir(parents=True, exist_ok=True)

    chroma_client = PersistentClient(
        path=str(permanent_chroma_path_text),
        settings=Settings(anonymized_telemetry=False)
    )
    collection = chroma_client.get_or_create_collection(
        name="default", metadata={"hnsw:space": "cosine"}
    )
    collection.add(
        documents=valid_texts,
        embeddings=embeddings_list,
        metadatas=valid_metadatas,
        ids=[str(uuid4()) for _ in valid_texts]
    )

    vectordb = Chroma(
        client=chroma_client,
        collection_name="default",
        persist_directory=str(permanent_chroma_path_text)
    )

    if "text_vector_db" in st.session_state:
        del st.session_state["text_vector_db"]
    

    elapsed_time = time.time() - start_time
    print(f"✅ Text documents stored in DB in {elapsed_time:.2f} seconds")

    return vectordb

text_db_path = ROOT_FOLDER / "chroma_text_db"


# Track processed .md file paths
processed_notes_log = ROOT_FOLDER / "processed_notes.txt"
if not processed_notes_log.exists():
    processed_notes_log.write_text("")  # create empty log if not present

processed_notes = set(processed_notes_log.read_text().splitlines())
current_notes = set(str(p) for p in NOTES_FOLDER.glob("*.md"))
new_notes = current_notes - processed_notes

        

if "text_vector_db" not in st.session_state:
    if text_db_path.exists():
        chroma_client = PersistentClient(
            path=str(text_db_path),
            settings=Settings(anonymized_telemetry=False)
        )
        vectordb = Chroma(
            client=chroma_client,
            collection_name="default",
            persist_directory=str(text_db_path)
        )

        # Check for new .md or .pdf files not in DB
        try:
            existing_sources = set(d.get("source") for d in vectordb.get()["metadatas"] if "source" in d)
        except:
            existing_sources = set()

        unprocessed_docs = [doc for doc in UPLOAD_FOLDER.glob("*") 
                    if doc.suffix.lower() in [".pdf", ".docx"] and str(doc) not in existing_sources]
        unprocessed_notes = [Path(p) for p in new_notes if p not in existing_sources]
        new_docs = unprocessed_docs  + unprocessed_notes

        if new_docs:
            with st.spinner("🔄 Adding new documents to DB..."):
                new_splits = load_docs(new_docs)
                for doc in new_splits:
                    emb = get_embedding(doc.page_content)
                    if emb and any(emb):
                        vectordb._collection.add(
                            documents=[doc.page_content],
                            embeddings=[emb],
                            metadatas=[doc.metadata],
                            ids=[str(uuid4())]
                        )

                # ✅ Log newly processed .md files
                new_md_sources = [str(doc.metadata["source"]) for doc in new_splits if doc.metadata["source"].endswith(".md")]
                updated_notes = processed_notes.union(new_md_sources)
                processed_notes_log.write_text("\n".join(updated_notes))
                        
                        
            print("✅ New documents added to DB")

        st.session_state.text_vector_db = vectordb
        print("✅ Loaded existing text DB")
        
        
        

    else:
        with st.spinner("Setting up text document DB..."):
            doc_splits = load_docs(document_paths)
            if doc_splits:
                st.session_state.text_vector_db = create_db(doc_splits)
            else:
                st.warning("⚠️ No valid documents found. Skipping DB creation.")

text_vector_db = st.session_state.text_vector_db

if st.session_state.text_vector_db:
    text_vector_db = st.session_state.text_vector_db
    # safe to use
else:
    st.info("ℹ️ Text DB not available. Upload documents to build it.")


image_output_folder = ROOT_FOLDER / "image_data"
image_output_folder.mkdir(parents=True, exist_ok=True)



def extract_images_and_metadata(document_paths):
    start_time = time.time()
    image_metadata = []

    for file_path in tqdm(document_paths, desc="Processing Documents"):
        doc_title = file_path.stem

        # --- 📄 PDF Handling ---
        if file_path.suffix.lower() == ".pdf":
            doc = fitz.open(file_path)

            for page_number in range(len(doc)):
                page = doc[page_number]
                image_list = page.get_images(full=True)

                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    image = Image.open(BytesIO(image_bytes)).convert("RGB")

                    image_filename = f"{doc_title}_page{page_number+1}_img{img_index+1}.{image_ext}"
                    image_path = image_output_folder / image_filename
                    image.save(image_path)

                    text = page.get_text("text")
                    lines = text.split("\n")
                    figure_captions = [
                        line.strip() for line in lines
                        if "figure" in line.lower() or "fig." in line.lower()
                    ]

                    caption = figure_captions[img_index] if img_index < len(figure_captions) else "No caption found"

                    metadata_entry = {
                        "filename": str(image_path),
                        "page_number": page_number + 1,
                        "caption": caption,
                        "source": str(file_path),
                        "document_title": doc_title,
                        "description": caption  
                    }
                    image_metadata.append(metadata_entry)

            doc.close()

        # --- 📝 DOCX Handling ---
        elif file_path.suffix.lower() == ".docx":
            try:
                doc = DocxDocument(file_path)
                rels = doc.part._rels

                for i, rel in enumerate(rels):
                    rel_obj = rels[rel]
                    if "image" in rel_obj.target_ref:
                        img_data = rel_obj.target_part.blob
                        image = Image.open(BytesIO(img_data)).convert("RGB")
                        image_filename = f"{doc_title}_img{i+1}.png"
                        image_path = image_output_folder / image_filename
                        image.save(image_path)

                        # 🧠 Try to extract caption from nearby paragraph
                        caption = f"Image {i+1}"  # default
                        for j, para in enumerate(doc.paragraphs):
                            if rel_obj.rId in para._p.xml:
                                # Look at next paragraph for caption
                                if j + 1 < len(doc.paragraphs):
                                    next_para = doc.paragraphs[j + 1].text.strip()
                                    if any(tag in next_para.lower() for tag in ["figure", "fig.", "image"]):
                                        caption = next_para
                                break

                        metadata_entry = {
                            "filename": str(image_path),
                            "page_number": i + 1,
                            "caption": caption,
                            "source": str(file_path),
                            "document_title": doc_title,
                            "description": caption
                        }
                        image_metadata.append(metadata_entry)

            except Exception as e:
                print(f"⚠️ Error reading images from {file_path.name}: {e}")

    # --- Final assembly ---
    image_docs = []
    for item in image_metadata:
        caption_text = item.get("caption", "")
        metadata = {
            "filename": item["filename"],
            "page_number": item["page_number"],
            "source": item["source"],
            "document_title": item["document_title"]
        }
        image_docs.append(Document(page_content=caption_text, metadata=metadata))

    print(f"\n✅ Extracted {len(image_metadata)} images.")
    elapsed_time = time.time() - start_time
    print(f"✅ Image documents stored in DB in {elapsed_time:.2f} seconds")

    return image_docs




def create_image_db(splits, db_folder_name):
    print("📸 Creating image DB with", len(splits), "docs")
    
    if not splits:
        print("⚠️ No image documents to process. Skipping image DB creation.")
        return None  # <- EARLY EXIT
    
    start_time = time.time()
    
    # Extract text and metadata
    texts = [doc.page_content for doc in splits]
    metadatas = [doc.metadata for doc in splits]

    # Manually get embeddings
    embeddings_list = []
    for text in texts:
        emb = get_embedding(text)
        if emb is not None:
            embeddings_list.append(emb)
        else:
            embeddings_list.append(np.zeros(384).tolist())

    # ✅ Create directly in permanent folder
    permanent_chroma_path_image = ROOT_FOLDER / "chroma_image_db"
    permanent_chroma_path_image.mkdir(parents=True, exist_ok=True)

    chroma_client = PersistentClient(
        path=str(permanent_chroma_path_image),
        settings=Settings(anonymized_telemetry=False)
    )
    collection = chroma_client.get_or_create_collection(
        name="default", metadata={"hnsw:space": "cosine"}
    )

    # Add documents and embeddings
    collection.add(
        documents=texts,
        embeddings=embeddings_list,
        metadatas=metadatas,
        ids=[str(uuid.uuid4()) for _ in texts]
    )

    vectordb = Chroma(
        client=chroma_client,
        collection_name="default",
        persist_directory=str(permanent_chroma_path_image)
    )

    # Clean up old state and rerun
    if "image_vector_db" in st.session_state:
        del st.session_state["image_vector_db"]
    # st.rerun()

    elapsed_time = time.time() - start_time
    print(f"✅ Image DB creation completed in {elapsed_time:.2f} seconds")

    return vectordb


image_db_path = ROOT_FOLDER / "chroma_image_db"

if "image_vector_db" not in st.session_state:
    if image_db_path.exists():
        chroma_client = PersistentClient(
            path=str(image_db_path),
            settings=Settings(anonymized_telemetry=False)
        )
        st.session_state.image_vector_db = Chroma(
            client=chroma_client,
            collection_name="default",
            persist_directory=str(image_db_path)
        )
        print("✅ Loaded existing image DB")
    else:
        with st.spinner("Setting up image document DB..."):
            image_docs = extract_images_and_metadata(document_paths)
            if image_docs:
                st.session_state.image_vector_db = create_image_db(image_docs, "chroma_image_db")
            else:
                st.warning("⚠️ No images found in uploaded PDFs. Skipping image DB creation.")

if st.session_state.image_vector_db:
    image_vector_db = st.session_state.image_vector_db
    # safe to use
else:
    st.info("ℹ️ Image DB not available. Upload documents to build it.")



def extract_tables_and_metadata(document_paths):
    start_time = time.time()
    table_docs = []
    table_count = 0

    for file_path in tqdm(document_paths, desc="Extracting Tables"):
        doc_title = file_path.stem

        # --- 📄 PDF Handling ---
        if file_path.suffix.lower() == ".pdf":
            try:
                tables = camelot.read_pdf(str(file_path), pages="all", flavor="lattice", strip_text="\n")
                if tables.n == 0:
                    tables = camelot.read_pdf(str(file_path), pages="all", flavor="stream", strip_text="\n")

                combined_df = pd.concat([t.df for t in tables], ignore_index=True)

                html = combined_df.to_html(index=False)
                text = combined_df.to_string(index=False, header=False)

                table_count += 1
                metadata = {
                    "filename": str(file_path),
                    "page_number": "multiple",
                    "document_title": doc_title,
                    "caption": f"Extracted Table {table_count}",
                    "table_id": f"{doc_title}_table{table_count}",
                    "html": html
                }

                table_docs.append(Document(page_content=html, metadata=metadata))
            except Exception as e:
                print(f"⚠️ Error extracting tables from PDF {file_path.name}: {e}")

        # --- 📝 DOCX Handling ---
        elif file_path.suffix.lower() == ".docx":
            try:
                doc = DocxDocument(file_path)
                for i, table in enumerate(doc.tables):
                    data = []
                    for row in table.rows:
                        data.append([cell.text.strip() for cell in row.cells])
                    df = pd.DataFrame(data)

                    html = df.to_html(index=False)
                    text = df.to_string(index=False, header=False)

                    table_count += 1
                    metadata = {
                        "filename": str(file_path),
                        "page_number": f"table_{i+1}",
                        "document_title": doc_title,
                        "caption": f"Extracted Table {table_count}",
                        "table_id": f"{doc_title}_table{table_count}",
                        "html": html
                    }

                    table_docs.append(Document(page_content=html, metadata=metadata))
            except Exception as e:
                print(f"⚠️ Error extracting tables from DOCX {file_path.name}: {e}")

    elapsed_time = time.time() - start_time
    print(f"✅ Table documents stored in DB in {elapsed_time:.2f} seconds")

    return table_docs





def create_table_db(splits, db_folder_name):
    
    print("📋 Creating table DB with", len(splits), "docs")
    start_time = time.time()

    if not splits:
        print("⚠️ No table documents to process. Skipping table DB creation.")
        return None

    texts = [doc.page_content for doc in splits]
    metadatas = [doc.metadata for doc in splits]

    embeddings_list = []
    for text in texts:
        emb = get_embedding(text)
        if emb:
            embeddings_list.append(emb)
        else:
            embeddings_list.append(np.zeros(384).tolist())

    if not any(embeddings_list):
        raise ValueError("❌ No valid embeddings for table documents. Cannot create DB.")

    permanent_chroma_path_table = ROOT_FOLDER / "chroma_table_db"
    permanent_chroma_path_table.mkdir(parents=True, exist_ok=True)

    chroma_client = PersistentClient(
        path=str(permanent_chroma_path_table),
        settings=Settings(anonymized_telemetry=False)
    )
    collection = chroma_client.get_or_create_collection(
        name="default", metadata={"hnsw:space": "cosine"}
    )

    collection.add(
        documents=texts,
        embeddings=embeddings_list,
        metadatas=metadatas,
        ids=[str(uuid4()) for _ in texts]
    )

    vectordb = Chroma(
        client=chroma_client,
        collection_name="default",
        persist_directory=str(permanent_chroma_path_table)
    )

    if "table_vector_db" in st.session_state:
        del st.session_state["table_vector_db"]
    # st.rerun()

    elapsed_time = time.time() - start_time
    print(f"✅ Table DB creation completed in {elapsed_time:.2f} seconds")

    return vectordb



table_db_path = ROOT_FOLDER / "chroma_table_db"

if "table_vector_db" not in st.session_state:
    if table_db_path.exists():
        chroma_client = PersistentClient(
            path=str(table_db_path),
            settings=Settings(anonymized_telemetry=False)
        )
        st.session_state.table_vector_db = Chroma(
            client=chroma_client,
            collection_name="default",
            persist_directory=str(table_db_path)
        )
        print("✅ Loaded existing table DB")
    else:
        with st.spinner("Setting up table document DB..."):
            table_docs = extract_tables_and_metadata(document_paths)
            if table_docs:
                table_db = create_table_db(table_docs, "chroma_table_db")
                if table_db:
                    st.session_state.table_vector_db = table_db
            else:
                warn = st.empty()
                warn.warning("⚠️ No table data found in uploaded documents. Skipping table DB creation.")
                time.sleep(3)
                warn.empty()

            

# table_vector_db = st.session_state.table_vector_db
table_vector_db = st.session_state.get("table_vector_db")





global_history = []

def conversation(qa_chain, message):
    global global_history

    if "text_vector_db" not in st.session_state:
        st.info("ℹ️ Vector DB not initialized. Please upload documents or process notes.")
        return "Vector DB is not initialized yet. Upload PDFs or process notes first.", [], [], []

    text_vector_db = st.session_state.get("text_vector_db", None)
    if text_vector_db is None:
        return "Vector DB is not initialized yet. Upload PDFs or process notes first.", [], [], []

    # Retrieve memory-based chat history
    memory_variables = memory.load_memory_variables({})
    chat_history = memory_variables.get('history', '')

    messages = [{"role": "system", "content": system_prompt + f"\n\nChat History:\n{chat_history}"}]
    messages.append({"role": "user", "content": message})

    # Retrieve top 3 relevant text documents
    query_embedding = get_embedding(message)
    selected = st.session_state.get("selected_docs", set())
    if selected:
        all_docs = text_vector_db.similarity_search_by_vector(query_embedding, k=10)
        similar_docs = [doc for doc in all_docs if os.path.basename(doc.metadata.get("source", "")) in selected][:3]
    else:
        similar_docs = []

    context = "\n\n".join(doc.page_content for doc in similar_docs)

    # Build LLaMA input prompt
    full_query = f"""You must answer based on the following documents:

{context}

User Query: {message}"""

    # Query LLaMA (or GPT if fallback)
    response_answer = process_query(full_query)

    # Save memory history
    memory.save_context({"input": message}, {"output": response_answer})

    # Extract text doc references
    references = set(
        os.path.basename(doc.metadata["source"])
        for doc in similar_docs
        if "source" in doc.metadata
    )

    # --- Image Retrieval ---
    image_outputs = []
    image_vector_db = st.session_state.get("image_vector_db")
    if image_vector_db:
        image_query_emb = get_embedding(message)
        image_results = image_vector_db.similarity_search_by_vector(image_query_emb, k=2)
        for img_doc in image_results:
            meta = img_doc.metadata
            filename = meta.get("filename", "")
            caption = img_doc.page_content
            doc_title = os.path.basename(meta.get("source", meta.get("document_title", "Unknown Title")))
            sim_score = cosine_similarity(image_query_emb, get_embedding(caption))
            image_outputs.append({
                "filename": filename,
                "caption": caption,
                "title": doc_title,
                "page": meta.get("page_number", "?"),
                "relevance": sim_score
            })

    # --- Table Retrieval ---
    table_outputs = []
    table_vector_db = st.session_state.get("table_vector_db")
    if should_display_table(message) and table_vector_db:
        table_query_emb = get_embedding(message)
        table_results = table_vector_db.similarity_search_by_vector(table_query_emb, k=2)
        for tbl_doc in table_results:
            meta = tbl_doc.metadata
            caption = meta.get("caption", "")
            doc_title = os.path.basename(meta.get("filename", meta.get("document_title", "Unknown Title")))
            sim_score = cosine_similarity(table_query_emb, get_embedding(caption))
            table_outputs.append({
                "caption": caption,
                "content": tbl_doc.page_content,
                "title": doc_title,
                "page": meta.get("page_number", "?"),
                "relevance": sim_score
            })

    # Manage local history buffer
    new_history_entry = (message, response_answer)
    global_history.append(new_history_entry)
    if len(global_history) > 10:
        global_history.pop(0)

    return response_answer, list(references), image_outputs, table_outputs


image_display_keywords = ["show", "display", "image", "figure", "chart", "plot"]

def should_display_image(user_input):
    return any(keyword in user_input.lower() for keyword in image_display_keywords)


table_display_keywords = ["table", "tabular", "show table", "display table", "tabular data"]

def should_display_table(user_input):
    return any(keyword in user_input.lower() for keyword in table_display_keywords)

def load_chat_to_memory(file_path, memory):
    with open(file_path) as f:
        data = json.load(f)
        for msg in data["messages"]:
            role = msg["role"]
            content = msg["content"]
            if role == "user":
                memory.chat_memory.add_message(HumanMessage(content=content))
            elif role == "assistant":
                memory.chat_memory.add_message(AIMessage(content=content))
            
def generate_auto_title(first_message):
    """Generate a clean title using GPT summarization of the first user message."""
    try:
        # Use your GPT endpoint to summarize or rephrase the first message
        summary_prompt = [
            {"role": "system", "content": "Summarize this user query into a short 3-5 word title suitable for a chat label. Remove punctuation."},
            {"role": "user", "content": first_message}
        ]
        title = query_llama(summary_prompt).strip()
        title = title.replace(" ", "_")[:30]
        return title or f"Chat_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}"
    except Exception as e:
        print(f"Failed to summarize title: {e}")
        return f"Chat_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}"



def show_menu_in_jupyter(text):
    display(Markdown(f"```\n{text}\n```"))


def select_chat():
    global active_chat_file, active_chat_title
    chats = list_previous_chats()
    menu_text = "Choose a conversation:\n"
    menu_text += "1. ➕ Start a new conversation\n"
    for idx, chat in enumerate(chats, start=2):
        _, file, title, turns = chat
        menu_text += f"{idx}. {title} ({turns} turns)\n"

    show_menu_in_jupyter(menu_text)

    choice = input("Select (enter number): ")
    try:
        choice = int(choice)
        if choice == 1:
            first_message = input("You: ")
            active_chat_title = generate_auto_title(first_message)
            active_chat_file = CHAT_FOLDER / f"chat_{active_chat_title}.json"
            return first_message  # important: return first message to immediately send
        elif 2 <= choice <= len(chats) + 1:
            selected_file = chats[choice - 2][1]
            active_chat_file = selected_file
            active_chat_title = selected_file.stem.replace("chat_", "")
            load_chat_to_memory(selected_file, memory)
            print(f"\n✅ Loaded previous conversation: {active_chat_title}")
            with open(selected_file) as f:
                data = json.load(f)
                for msg in data["messages"]:
                    role = msg["role"]
                    content = msg["content"]
                    print(f"{role.capitalize()}: {content}")
            return None
    except ValueError:
        print("Invalid input, starting a new conversation.")
        return None
    
    
# --- STREAMLIT CHAT UI ---

# --- Session state init ---
defaults = {
    "active_chat_title": "",
    "active_chat_file": None,
    "chat_history": [],
    "chat_saved": False,
    "rename_target": None,
    "rename_active": False
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# --- SIDEBAR: Start or Load Chat ---
with st.sidebar:

    # --- Process Notes ---
    if new_notes:
        if st.button("📄 Process new notes from Notes folder"):
            with st.spinner("Processing new notes..."):
                new_note_docs = load_docs([Path(p) for p in new_notes])

                if not new_note_docs:
                    st.warning("⚠️ No valid notes found to process.")
                    st.stop()

                if "text_vector_db" in st.session_state and st.session_state.text_vector_db:
                    vectordb = st.session_state.text_vector_db
                    for doc in new_note_docs:
                        emb = get_embedding(doc.page_content)
                        if emb and any(emb):
                            vectordb._collection.add(
                                documents=[doc.page_content],
                                embeddings=[emb],
                                metadatas=[doc.metadata],
                                ids=[str(uuid4())]
                            )
                else:
                    st.session_state.text_vector_db = create_db(new_note_docs)

                processed_notes_log.write_text("\n".join(current_notes))
                st.success("✅ New notes added to DB.")
                st.rerun()

    # --- Upload PDFs ---
    st.markdown("### 📤 Upload PDFs")

    if "uploader_key" not in st.session_state:
        st.session_state["uploader_key"] = 0
        
    if "new_uploads" not in st.session_state:
        st.session_state["new_uploads"] = []
        
    if "selected_docs" not in st.session_state:
        st.session_state.selected_docs = set()

    uploaded_files = st.file_uploader(
        "Upload one or more Docs", 
        type=["pdf", "docx"], 
        accept_multiple_files=True,
        key=f"pdf_upload_{st.session_state['uploader_key']}"
    )

    if uploaded_files:
        uploaded_names = []
        for uploaded_file in uploaded_files:
            if uploaded_file.size > 200 * 1024 * 1024:
                st.warning(f"❌ {uploaded_file.name} exceeds 200MB limit.")
                continue
            file_path = UPLOAD_FOLDER / uploaded_file.name
            if not file_path.exists():
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.read())
                st.session_state["new_uploads"].append(file_path)
                uploaded_names.append(uploaded_file.name)
            else:
                uploaded_names.append(uploaded_file.name)

        if uploaded_names:
            st.success(f"✅ Uploaded: {', '.join(uploaded_names)}")

        # Always show Process button
        if st.button("📄 Process Uploaded Docs"):
            if "text_vector_db" in st.session_state and st.session_state.text_vector_db:
                existing_sources = set(d.get("source") for d in st.session_state.text_vector_db.get()["metadatas"])
            else:
                existing_sources = set()

            to_process = []
            skipped_files = []

            for file_path in st.session_state["new_uploads"]:
                if str(file_path) in existing_sources:
                    skipped_files.append(file_path.name)
                else:
                    to_process.append(file_path)

            if skipped_files:
                st.warning(f"⚠️ Skipping already processed: {', '.join(skipped_files)}")

            if to_process:
                with st.spinner("🔄 Processing uploaded documents..."):
                    doc_splits = load_docs(to_process)

                    if "text_vector_db" in st.session_state and st.session_state.text_vector_db:
                        vectordb = st.session_state.text_vector_db
                        for doc in doc_splits:
                            emb = get_embedding(doc.page_content)
                            if emb and any(emb):
                                vectordb._collection.add(
                                    documents=[doc.page_content],
                                    embeddings=[emb],
                                    metadatas=[doc.metadata],
                                    ids=[str(uuid4())]
                                )
                    else:
                        st.session_state.text_vector_db = create_db(doc_splits)

                    # ✅ Image + Table
                    image_docs = extract_images_and_metadata(to_process)
                    if image_docs:
                        if "image_vector_db" in st.session_state and st.session_state.image_vector_db:
                            for doc in image_docs:
                                emb = get_embedding(doc.page_content)
                                if emb:
                                    st.session_state.image_vector_db._collection.add(
                                        documents=[doc.page_content],
                                        embeddings=[emb],
                                        metadatas=[doc.metadata],
                                        ids=[str(uuid4())]
                                    )
                        else:
                            st.session_state.image_vector_db = create_image_db(image_docs, "chroma_image_db")

                    table_docs = extract_tables_and_metadata(to_process)
                    if table_docs:
                        if "table_vector_db" in st.session_state and st.session_state.table_vector_db:
                            for doc in table_docs:
                                emb = get_embedding(doc.page_content)
                                if emb:
                                    st.session_state.table_vector_db._collection.add(
                                        documents=[doc.page_content],
                                        embeddings=[emb],
                                        metadatas=[doc.metadata],
                                        ids=[str(uuid4())]
                                    )
                        else:
                            st.session_state.table_vector_db = create_table_db(table_docs, "chroma_table_db")

                    st.success("✅ Processing completed.")
            else:
                st.info("ℹ️ No new documents to process.")

            # Clean up
            st.session_state["uploader_key"] += 1
            st.session_state["new_uploads"].clear()
            time.sleep(1)
            st.rerun()

            

    # --- Document List ---
    st.markdown("### 📚 Documents in DB")
    if "text_vector_db" in st.session_state and st.session_state.text_vector_db:
        vectordb = st.session_state.text_vector_db
        try:
            sources = sorted(set(doc.get("source", "Unknown") for doc in vectordb.get()["metadatas"]))
            with st.expander("📄 View processed documents", expanded=False):
                total_docs = len(sources)
                page_size = 5
                max_page = (total_docs - 1) // page_size + 1

                if "doc_page" not in st.session_state:
                    st.session_state.doc_page = 1

                col1, col2, col3 = st.columns([1,2,1])

                with col1:
                    if st.button("⬅️") and st.session_state.doc_page > 1:
                        st.session_state.doc_page -= 1
                with col3:
                    if st.button("➡️") and st.session_state.doc_page < max_page:
                        st.session_state.doc_page += 1
                with col2:
                    st.write(f"Page {st.session_state.doc_page} of {max_page}")

                start_idx = (st.session_state.doc_page - 1) * page_size
                end_idx = min(start_idx + page_size, total_docs)

                shown_sources = list(sources)[start_idx:end_idx]

                for src in shown_sources:
                    file_name = os.path.basename(src)
                    checked = file_name in st.session_state.selected_docs

                    col_doc, col_del = st.columns([6, 1])  # wider checkbox, small delete

                    with col_doc:
                        if st.checkbox(file_name, value=checked, key=f"chk_{file_name}"):
                            st.session_state.selected_docs.add(file_name)
                        else:
                            st.session_state.selected_docs.discard(file_name)

                    with col_del:
                        if st.button("🗑", key=f"del_{file_name}"):
                            try:
                                if "text_vector_db" in st.session_state:
                                    vectordb = st.session_state.text_vector_db
                                    all_meta = vectordb.get()
                                    ids_to_delete = [
                                        id_
                                        for id_, meta in zip(all_meta["ids"], all_meta["metadatas"])
                                        if meta.get("source", "") == src
                                    ]
                                    vectordb._collection.delete(ids=ids_to_delete)
                                    st.success(f"Deleted")

                                    # optional: also remove physical file
                                    file_path = Path(src)
                                    if file_path.exists():
                                        file_path.unlink()
                                        st.info(f"🗑 File {file_name} removed from disk")

                                    st.rerun()
                                else:
                                    st.warning("Vector DB not loaded")
                            except Exception as e:
                                st.error(f"❌ Could not delete: {e}")

        except Exception as e:
            st.warning(f"⚠️ Could not load document list. Error: {e}")

    # --- Chat Section ---
    st.header("💬 Start / Load Chat")

    if st.button("➕ Start New Chat"):
        memory.clear()
        st.session_state.chat_history = []
        st.session_state.active_chat_file = None
        st.session_state.active_chat_title = None
        st.session_state.chat_saved = False
        st.rerun()

    st.markdown("### 📂 Previous Chats")
    chat_files = sorted(CHAT_FOLDER.glob("chat_*.json"), key=os.path.getmtime, reverse=True)[:10]

    for file in chat_files:
        title_raw = file.stem.replace("chat_", "")
        title = title_raw.replace("_", " ")
        with st.container():
            col1, col2 = st.columns([8, 1.5])
            with col1:
                if title_raw == st.session_state.active_chat_title:
                    st.markdown(f"**📄 {title}**")
                else:
                    if st.button(f"📄 {title}", key=f"load_{title}"):
                        memory.clear()
                        load_chat_to_memory(file, memory)
                        st.session_state.active_chat_title = title_raw
                        st.session_state.active_chat_file = file
                        with open(file) as f:
                            data = json.load(f)
                            st.session_state.chat_history = data.get("messages", [])
                        st.rerun()
            with col2:
                menu_key = f"menu_{title}"
                if st.button("⋮", key=menu_key):
                    st.session_state[f"show_menu_{menu_key}"] = not st.session_state.get(f"show_menu_{menu_key}", False)

            if st.session_state.get(f"show_menu_{menu_key}", False):
                with st.popover(f"⚙️ Options for {title}"):
                    if st.button("📝 Rename", key=f"rename_{title}"):
                        st.session_state.rename_target = file
                        st.session_state.rename_active = True
                    if st.button("🗑 Delete", key=f"delete_{title}"):
                        file.unlink()
                        st.session_state.pop(f"show_menu_{menu_key}", None)
                        st.rerun()

    # --- Rename Chat ---
    if st.session_state.get("rename_active", False):
        new_name = st.text_input("Enter new title", value=st.session_state.rename_target.stem.replace("chat_", ""))
        if st.button("✅ Rename"):
            old_path = st.session_state.rename_target
            new_title = new_name.replace(" ", "_")
            new_path = CHAT_FOLDER / f"chat_{new_title}.json"
            old_path.rename(new_path)
            st.session_state.rename_active = False
            st.rerun()
        if st.button("❌ Cancel"):
            st.session_state.rename_active = False

# --- Chat History Display ---
for i, msg in enumerate(st.session_state.chat_history):
    role = msg["role"]
    content = msg["content"]

    with st.chat_message(role):
        st.markdown(content)

        if role == "assistant":
            # ✅ Add Read Aloud button with unique key per message
            speak_text_button(content, key_suffix=f"history_{i}")

            # ✅ Show references only if not embedded
            if "references" in msg and msg["references"] and "[SOURCE:" not in content:
                st.markdown("**References:**")
                for ref in msg["references"]:
                    st.markdown(f"- {ref}")

            # ✅ Show images (if any)
            if "images" in msg and msg["images"]:
                for img in msg["images"]:
                    st.image(
                        img["filename"],
                        caption=f"{img['caption']} ({img['title']}, Page {img['page']})",
                        width=300
                    )

# --- Chat Input ---
user_input = st.chat_input("Ask a question about the documents...")

if "pending_input" in st.session_state:
    user_input = st.session_state.pop("pending_input")
    
if user_input:
    with st.chat_message("user"):
        st.markdown(user_input)
        
    if not st.session_state.selected_docs:
        st.warning("⚠️ No documents selected. Please select at least one document from the sidebar to answer questions.")
        st.stop()

    if user_input.strip().lower() == "quit":
        active_file = st.session_state.get("active_chat_file")
        if active_file:
            messages = []
            for msg in st.session_state.chat_history:
                entry = {"role": msg["role"], "content": msg["content"]}
                if "images" in msg and msg["images"]:
                    entry["images"] = msg["images"]
                messages.append(entry)

            with open(active_file, "w") as f:
                json.dump({
                    "title": st.session_state.active_chat_title,
                    "messages": messages
                }, f)
            st.success("✅ Chat saved!")
        else:
            st.warning("⚠️ No active chat file found. Chat not saved.")
        st.stop()

    if len(st.session_state.chat_history) == 0:
        summarized_title = generate_auto_title(user_input)
        new_file = CHAT_FOLDER / f"chat_{summarized_title}.json"
        old_file = st.session_state.active_chat_file

        if old_file and isinstance(old_file, Path) and old_file.exists():
            try:
                old_file.rename(new_file)
            except Exception as e:
                print(f"Rename failed: {e}")
        

        st.session_state.active_chat_file = new_file
        st.session_state.active_chat_title = summarized_title
        st.session_state.chat_saved = True
        
    # --- Process and Answer ---
    # --- Process and Answer ---
    answer, references, image_outputs, table_outputs = conversation(None, user_input)

    # ✅ Append messages BEFORE saving
    st.session_state.chat_history.append({
        "role": "user",
        "content": user_input
    })
    st.session_state.chat_history.append({
    "role": "assistant",
    "content": answer,
    "images": image_outputs if should_display_image(user_input) else None,
    "references": references
})
    
    
    # ✅ Generate file and title if not already set
    if st.session_state.active_chat_file is None:
        summarized_title = generate_auto_title(user_input)
        st.session_state.active_chat_title = summarized_title
        st.session_state.active_chat_file = CHAT_FOLDER / f"chat_{summarized_title}.json"

    # ✅ Save the chat to JSON
    active_file = st.session_state.active_chat_file
    if active_file and isinstance(active_file, Path):
        messages = []
        for msg in st.session_state.chat_history:
            entry = {
                "role": msg["role"],
                "content": msg["content"]
            }
            if "images" in msg and msg["images"]:
                entry["images"] = msg["images"]
            if "references" in msg and msg["references"]:
                entry["references"] = msg["references"]
            messages.append(entry)

        try:
            with open(active_file, "w", encoding="utf-8") as f:
                json.dump({
                    "title": st.session_state.active_chat_title,
                    "messages": messages
                }, f, indent=2)
            print(f"✅ Chat saved to {active_file.name}")
        except Exception as e:
            st.error(f"❌ Could not save chat: {e}")
    else:
        st.warning("⚠️ Chat file path not set. Unable to save chat.")

    with st.chat_message("assistant"):
        stream_text_to_ui(answer)

        if references and "[SOURCE:" not in answer:
            st.markdown("**References:**")
            for ref in references:
                st.markdown(f"- {ref}")
        else:
            st.markdown("_No references found._")

        if should_display_image(user_input) or any(img["relevance"] > 0.6 for img in image_outputs):
            st.markdown("### 🎼 Relevant Images")
            for img in image_outputs:
                st.image(img["filename"], caption=f"{img['caption']} ({img['title']}, Page {img['page']})", width=300)

        if should_display_table(user_input) or any(tbl["relevance"] > 0.6 for tbl in table_outputs):
            st.markdown("### 📋 Relevant Tables")
            for tbl in table_outputs:
                st.markdown(f"**{tbl['caption']}** ({tbl['title']} - Page {tbl['page']})")
                components.html(tbl['content'], height=400, scrolling=True)
                
        speak_text_button(answer, key_suffix=str(len(st.session_state.chat_history)))

    
    

    # ✅ Rerun once to refresh sidebar after first message
    if not st.session_state.chat_saved:
        st.session_state.chat_saved = True
        st.rerun()